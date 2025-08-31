import streamlit as st
import PyPDF2
import docx
import io
from typing import Optional

class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    def extract_text(self, uploaded_file) -> str:
        """Extract text from uploaded file based on its type."""
        try:
            file_type = uploaded_file.type
            
            if file_type == "application/pdf":
                return self._extract_from_pdf(uploaded_file)
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                return self._extract_from_docx(uploaded_file)
            elif file_type == "text/plain":
                return self._extract_from_txt(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            st.error(f"Error extracting text: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file."""
        try:
            # Read the PDF file
            pdf_bytes = uploaded_file.read()
            pdf_file = io.BytesIO(pdf_bytes)
            
            # Create PDF reader object
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file."""
        try:
            # Read the DOCX file
            docx_bytes = uploaded_file.read()
            docx_file = io.BytesIO(docx_bytes)
            
            # Create document object
            doc = docx.Document(docx_file)
            
            # Extract text from all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, uploaded_file) -> str:
        """Extract text from TXT file."""
        try:
            # Read the text file
            content = uploaded_file.read()
            
            # Decode bytes to string
            if isinstance(content, bytes):
                # Try different encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        text = content.decode(encoding)
                        return text.strip()
                    except UnicodeDecodeError:
                        continue
                
                # If all encodings fail, use utf-8 with error handling
                text = content.decode('utf-8', errors='replace')
                return text.strip()
            else:
                return str(content).strip()
                
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    def get_document_info(self, uploaded_file) -> dict:
        """Get basic information about the uploaded document."""
        return {
            "filename": uploaded_file.name,
            "file_type": uploaded_file.type,
            "file_size": uploaded_file.size,
            "size_mb": round(uploaded_file.size / (1024 * 1024), 2)
        }
    
    def validate_file(self, uploaded_file) -> tuple[bool, str]:
        """Validate if the uploaded file is supported and within size limits."""
        max_size_mb = 50  # 50 MB limit
        supported_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "text/plain"
        ]
        
        # Check file type
        if uploaded_file.type not in supported_types:
            return False, f"Unsupported file type: {uploaded_file.type}"
        
        # Check file size
        size_mb = uploaded_file.size / (1024 * 1024)
        if size_mb > max_size_mb:
            return False, f"File size ({size_mb:.1f} MB) exceeds limit of {max_size_mb} MB"
        
        return True, "File is valid"
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Remove common PDF artifacts
        text = text.replace("\x00", "")  # Null characters
        text = text.replace("\ufffd", "")  # Replacement characters
        
        # Basic cleaning
        text = text.strip()
        
        return text
